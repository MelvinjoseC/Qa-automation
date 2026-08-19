import logging
from docx import Document
from exceptions import MDRParsingError

def parse_mdr_docx(mdr_path: str):
    """
    Parse the MDR .docx and return:
      required_folders: set of normalized relative folder paths
      required_files: set of normalized relative file paths
    """
    logging.info(f"Parsing MDR file: {mdr_path}")
    try:
        doc = Document(mdr_path)
    except Exception as e:
        raise MDRParsingError(f"Failed to open/parse MDR document: {e}") from e


    required_folders = set()
    required_files = set()

    # Parse regular paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Normalize slashes and strip leading "./"
        norm = text.replace("\\", "/").lstrip("./")

        if norm.endswith("/"):
            required_folders.add(norm.rstrip("/"))
        else:
            required_files.add(norm)

    # Parse tables (MDRs are often structured in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    norm = text.replace("\\", "/").lstrip("./")
                    if norm.endswith("/"):
                        required_folders.add(norm.rstrip("/"))
                    else:
                        required_files.add(norm)

    logging.info(f"MDR parse result: {len(required_folders)} folders, {len(required_files)} files.")
    return required_folders, required_files
