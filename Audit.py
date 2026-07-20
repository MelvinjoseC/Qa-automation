import os
import logging
from datetime import datetime
from pathlib import Path
from tkinter import Tk, Text, END, Scrollbar, RIGHT, Y, LEFT, BOTH, X, TOP, BOTTOM, Frame
from tkinter import ttk
import tkinter as tk

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

# =========================
# CONFIG / BRANDING
# =========================

COMPANY_NAME = "FUSIE Engineers"
COMPANY_TAGLINE = "Precision. Safety. Compliance."
COMPANY_LOGO_PATH = "branding/fusie_logo.png"  # put your logo here (optional)

LOG_DIR = "logs"
os.makedirs(LOG_DIR, exist_ok=True)
LOG_FILE = os.path.join(LOG_DIR, "iso_audit.log")

logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)

# =========================
# MDR PARSING ASSUMPTIONS
# =========================
"""
MDR FORMAT ASSUMPTION (simple but effective):

We assume the MDR .docx has one required item per line/paragraph, like:

  Project/
  Project/01_Management/
  Project/01_Management/QM-001 Quality Manual.docx
  Project/02_Design/
  Project/02_Design/DS-010 General Arrangement.dwg

Rules:
- If line ends with "/" or "\" → treat as REQUIRED FOLDER
- Else → treat as REQUIRED FILE
- Paths are treated as relative logical paths. We will normalize them.

You can later customize `parse_mdr_docx()` to match your company’s MDR style
(e.g. tables, clause columns, etc.).
"""


def parse_mdr_docx(mdr_path: str):
    """
    Parse the MDR .docx and return:
      required_folders: set of normalized relative folder paths
      required_files: set of normalized relative file paths
    """
    logging.info(f"Parsing MDR file: {mdr_path}")
    doc = Document(mdr_path)

    required_folders = set()
    required_files = set()

    # Parse regular paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Normalize slashes and strip leading "./"
        norm = text.replace("\\", "/").lstrip("./")

        if norm.endswith("/"):
            required_folders.add(norm.rstrip("/"))
        else:
            required_files.add(norm)

    # Parse tables (MDRs are often structured in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    norm = text.replace("\\", "/").lstrip("./")
                    if norm.endswith("/"):
                        required_folders.add(norm.rstrip("/"))
                    else:
                        required_files.add(norm)

    logging.info(f"MDR parse result: {len(required_folders)} folders, {len(required_files)} files.")
    return required_folders, required_files


def scan_project_structure(project_root: str):
    """
    Scan the actual folder structure and return:
      actual_folders: set of relative folder paths
      actual_files: set of relative file paths
    """
    logging.info(f"Scanning project folder: {project_root}")
    root_path = Path(project_root).resolve()
    actual_folders = set()
    actual_files = set()

    for dirpath, dirnames, filenames in os.walk(root_path):
        rel_dir = Path(dirpath).relative_to(root_path)
        if str(rel_dir) != ".":
            actual_folders.add(str(rel_dir).replace("\\", "/"))

        for f in filenames:
            file_rel_path = Path(dirpath).joinpath(f).relative_to(root_path)
            actual_files.add(str(file_rel_path).replace("\\", "/"))

    logging.info(f"Scan result: {len(actual_folders)} folders, {len(actual_files)} files.")
    return actual_folders, actual_files


def perform_gap_analysis(required_folders, required_files, actual_folders, actual_files, project_root=None):
    """
    Compare MDR requirements vs actual structure.

    Returns:
      nc_list: list of Non-Conformities
      obs_list: list of Observations
      ofi_list: list of Opportunities for Improvement
      summary: dict with counts
    """
    logging.info("Performing gap analysis...")

    nc_list = []
    obs_list = []
    ofi_list = []

    # Missing folders (NC)
    missing_folders = sorted(required_folders - actual_folders)
    for folder in missing_folders:
        nc_list.append({
            "type": "NC",
            "item_type": "Folder",
            "path": folder,
            "description": f"Required folder '{folder}' is missing.",
            "clause": "",
        })

    # Missing files (NC)
    missing_files = sorted(required_files - actual_files)
    for file in missing_files:
        nc_list.append({
            "type": "NC",
            "item_type": "File",
            "path": file,
            "description": f"Required document '{file}' is missing.",
            "clause": "",
        })

    # Extra folders (OBS)
    extra_folders = sorted(actual_folders - required_folders)
    for folder in extra_folders:
        obs_list.append({
            "type": "OBS",
            "item_type": "Folder",
            "path": folder,
            "description": f"Folder '{folder}' exists but is not defined in the MDR.",
            "clause": "",
        })

    # Extra files (OBS)
    extra_files = sorted(actual_files - required_files)
    for file in extra_files:
        obs_list.append({
            "type": "OBS",
            "item_type": "File",
            "path": file,
            "description": f"File '{file}' exists but is not defined in the MDR.",
            "clause": "",
        })

    # OFI check: Required folder exists but is empty
    if project_root:
        root_path = Path(project_root).resolve()
        for folder in sorted(required_folders & actual_folders):
            folder_path = root_path.joinpath(folder)
            if folder_path.exists() and folder_path.is_dir():
                try:
                    children = [c for c in folder_path.iterdir() if c.name not in (".DS_Store", "Thumbs.db")]
                    if not children:
                        ofi_list.append({
                            "type": "OFI",
                            "item_type": "Folder",
                            "path": folder,
                            "description": f"Required folder '{folder}' exists but is empty.",
                            "clause": "",
                        })
                except Exception as e:
                    logging.warning(f"Failed to check if folder {folder_path} is empty: {e}")

    summary = {
        "missing_folders": len(missing_folders),
        "missing_files": len(missing_files),
        "extra_folders": len(extra_folders),
        "extra_files": len(extra_files),
        "nc_count": len(nc_list),
        "obs_count": len(obs_list),
        "ofi_count": len(ofi_list),
    }

    logging.info(f"Gap analysis done: {summary}")
    return nc_list, obs_list, ofi_list, summary


def generate_pdf_report(
    output_path: str,
    project_root: str,
    mdr_path: str,
    required_folders,
    required_files,
    actual_folders,
    actual_files,
    nc_list,
    obs_list,
    ofi_list,
    summary,
):
    """
    Generate PDF report with branding using reportlab.
    """
    logging.info(f"Generating PDF report: {output_path}")

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=30,
        leftMargin=30,
        topMargin=40,
        bottomMargin=30,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='TitleCenter', alignment=1, fontSize=18, spaceAfter=12))
    styles.add(ParagraphStyle(name='SectionHeader', fontSize=14, spaceAfter=6, textColor=colors.HexColor("#003366")))
    styles.add(ParagraphStyle(name='NormalSmall', fontSize=9, spaceAfter=4))

    flow = []

    # Title / Branding
    title_text = f"{COMPANY_NAME} – ISO Project Folder Audit Report"
    flow.append(Paragraph(title_text, styles['TitleCenter']))
    flow.append(Paragraph(COMPANY_TAGLINE, styles['Normal']))
    flow.append(Spacer(1, 12))

    # Meta info
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    flow.append(Paragraph(f"<b>Audit Date:</b> {now_str}", styles['Normal']))
    flow.append(Paragraph(f"<b>Project Folder:</b> {project_root}", styles['Normal']))
    flow.append(Paragraph(f"<b>MDR Source:</b> {mdr_path}", styles['Normal']))
    flow.append(Spacer(1, 12))

    # Summary section
    flow.append(Paragraph("A. Summary", styles['SectionHeader']))
    summary_lines = [
        f"Missing Folders: {summary['missing_folders']}",
        f"Missing Files: {summary['missing_files']}",
        f"Extra Folders: {summary['extra_folders']}",
        f"Extra Files: {summary['extra_files']}",
        f"Total Non-Conformities (NC): {summary['nc_count']}",
        f"Total Observations (OBS): {summary['obs_count']}",
        f"Total Opportunities for Improvement (OFI): {summary['ofi_count']}",
    ]
    for line in summary_lines:
        flow.append(Paragraph(line, styles['NormalSmall']))
    flow.append(Spacer(1, 12))

    # MDR Requirements
    flow.append(Paragraph("B. MDR Requirements Overview", styles['SectionHeader']))
    flow.append(Paragraph(f"Defined Folders: {len(required_folders)}", styles['NormalSmall']))
    flow.append(Paragraph(f"Defined Documents: {len(required_files)}", styles['NormalSmall']))
    flow.append(Spacer(1, 6))

    # Actual Structure
    flow.append(Paragraph("C. Actual Project Folder Structure (Counts)", styles['SectionHeader']))
    flow.append(Paragraph(f"Detected Folders: {len(actual_folders)}", styles['NormalSmall']))
    flow.append(Paragraph(f"Detected Files: {len(actual_files)}", styles['NormalSmall']))
    flow.append(Spacer(1, 12))

    # Detailed NC Table
    flow.append(Paragraph("D. Non-Conformities (NC)", styles['SectionHeader']))
    if nc_list:
        nc_data = [["#", "Type", "Item Type", "Path", "Description"]]
        for i, nc in enumerate(nc_list, start=1):
            nc_data.append([
                str(i),
                nc["type"],
                nc["item_type"],
                nc["path"],
                nc["description"],
            ])
        nc_table = Table(nc_data, repeatRows=1, colWidths=[25, 35, 60, 200, 200])
        nc_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#003366")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ('FONTSIZE', (0, 1), (-1, -1), 7),
        ]))
        flow.append(nc_table)
    else:
        flow.append(Paragraph("No Non-Conformities detected.", styles['NormalSmall']))
    flow.append(Spacer(1, 12))

    # OBS
    flow.append(Paragraph("E. Observations (OBS)", styles['SectionHeader']))
    if obs_list:
        obs_data = [["#", "Type", "Item Type", "Path", "Description"]]
        for i, obs in enumerate(obs_list, start=1):
            obs_data.append([
                str(i),
                obs["type"],
                obs["item_type"],
                obs["path"],
                obs["description"],
            ])
        obs_table = Table(obs_data, repeatRows=1, colWidths=[25, 35, 60, 200, 200])
        obs_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#666666")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ('FONTSIZE', (0, 1), (-1, -1), 7),
        ]))
        flow.append(obs_table)
    else:
        flow.append(Paragraph("No Observations recorded.", styles['NormalSmall']))
    flow.append(Spacer(1, 12))

    # OFI
    flow.append(Paragraph("F. Opportunities for Improvement (OFI)", styles['SectionHeader']))
    if ofi_list:
        ofi_data = [["#", "Type", "Item Type", "Path", "Description"]]
        for i, ofi in enumerate(ofi_list, start=1):
            ofi_data.append([
                str(i),
                ofi["type"],
                ofi["item_type"],
                ofi["path"],
                ofi["description"],
            ])
        ofi_table = Table(ofi_data, repeatRows=1, colWidths=[25, 35, 60, 200, 200])
        ofi_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#006600")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ('FONTSIZE', (0, 1), (-1, -1), 7),
        ]))
        flow.append(ofi_table)
    else:
        flow.append(Paragraph("No Opportunities for Improvement identified.", styles['NormalSmall']))
    flow.append(Spacer(1, 12))

    # Corrective Action section (placeholder)
    flow.append(Paragraph("G. Corrective Action Summary", styles['SectionHeader']))
    flow.append(Paragraph(
        "For each NC, the responsible process owner shall define and implement corrective actions, "
        "including root cause analysis, target dates, and verification of effectiveness.",
        styles['NormalSmall']
    ))

    doc.build(flow)
    logging.info("PDF report generated successfully.")


# =========================
# TKINTER GUI
# =========================

class ISOAditorGUI:
    def __init__(self, master):
        self.master = master
        master.title(f"{COMPANY_NAME} - ISO Project Folder Auditor")
        master.configure(bg="#F4F6F9")

        self.mdr_path = None
        self.project_path = None

        # Setup modern styles
        style = ttk.Style()
        try:
            style.theme_use("clam")
        except Exception:
            pass

        primary_color = "#003366"  # Deep Navy
        secondary_color = "#336699" # Medium Blue
        bg_color = "#F4F6F9"
        accent_color = "#2E7D32"   # Forest Green
        text_dark = "#2D3748"

        style.configure(".", background=bg_color, font=("Segoe UI", 10), foreground=text_dark)
        style.configure("TFrame", background=bg_color)
        
        # Labels
        style.configure("Title.TLabel", font=("Segoe UI", 16, "bold"), foreground=primary_color, background=bg_color)
        style.configure("Subtitle.TLabel", font=("Segoe UI", 10, "italic"), foreground="#718096", background=bg_color)
        style.configure("Path.TLabel", font=("Segoe UI", 9), foreground="#4A5568", background=bg_color)
        style.configure("StatusRed.TLabel", font=("Segoe UI", 9, "bold"), foreground="#E53E3E", background=bg_color)
        style.configure("StatusGreen.TLabel", font=("Segoe UI", 9, "bold"), foreground="#38A169", background=bg_color)

        # Buttons
        style.configure("TButton", font=("Segoe UI", 10), padding=6)
        style.configure("Primary.TButton", font=("Segoe UI", 10, "bold"), foreground="white", background=primary_color)
        style.map("Primary.TButton", background=[("active", secondary_color)])
        
        style.configure("Action.TButton", font=("Segoe UI", 10, "bold"), foreground="white", background=accent_color)
        style.map("Action.TButton", background=[("active", "#22543D")])

        # Main Outer Container
        main_container = ttk.Frame(master, padding=15)
        main_container.pack(fill=BOTH, expand=True)

        # Top branding header
        header_frame = ttk.Frame(main_container)
        header_frame.pack(side=TOP, fill=X, pady=(0, 15))

        self.label_title = ttk.Label(
            header_frame,
            text=f"{COMPANY_NAME} - ISO Folder Audit Tool",
            style="Title.TLabel"
        )
        self.label_title.pack(side=TOP, anchor="w")

        self.label_sub = ttk.Label(
            header_frame,
            text=COMPANY_TAGLINE,
            style="Subtitle.TLabel"
        )
        self.label_sub.pack(side=TOP, anchor="w")

        # Divider
        separator = ttk.Separator(main_container, orient="horizontal")
        separator.pack(side=TOP, fill=X, pady=(0, 15))

        # File Selection Area (Card-like layout using LabelFrame)
        selection_frame = ttk.LabelFrame(main_container, text=" Configuration & Paths ", padding=12)
        selection_frame.pack(side=TOP, fill=X, pady=(0, 15))

        # MDR Document selection row
        mdr_row = ttk.Frame(selection_frame)
        mdr_row.pack(fill=X, pady=(0, 8))
        self.btn_mdr = ttk.Button(mdr_row, text="Browse MDR (.docx)", command=self.select_mdr, style="Primary.TButton")
        self.btn_mdr.pack(side=LEFT)
        
        self.label_mdr_status = ttk.Label(mdr_row, text="No MDR selected", style="StatusRed.TLabel", padding=(10, 0))
        self.label_mdr_status.pack(side=LEFT, fill=X, expand=True)

        # Project Folder selection row
        project_row = ttk.Frame(selection_frame)
        project_row.pack(fill=X, pady=(0, 8))
        self.btn_project = ttk.Button(project_row, text="Browse Project Folder", command=self.select_project_folder, style="Primary.TButton")
        self.btn_project.pack(side=LEFT)

        self.label_project_status = ttk.Label(project_row, text="No Project Folder selected", style="StatusRed.TLabel", padding=(10, 0))
        self.label_project_status.pack(side=LEFT, fill=X, expand=True)

        # Action Buttons frame
        actions_row = ttk.Frame(selection_frame)
        actions_row.pack(fill=X, pady=(8, 0))
        self.btn_run = ttk.Button(actions_row, text="Run Folder Audit & Generate PDF Report", command=self.run_audit, style="Action.TButton")
        self.btn_run.pack(side=LEFT, ipady=2)

        # Log Console Area
        console_frame = ttk.LabelFrame(main_container, text=" Audit Execution Console ", padding=10)
        console_frame.pack(side=TOP, fill=BOTH, expand=True)

        self.text_output = Text(
            console_frame,
            wrap="word",
            height=12,
            font=("Consolas", 9),
            bg="#1E293B",  # Dark slate background
            fg="#F8FAFC",  # Near white text
            insertbackground="white",
            relief="flat",
            borderwidth=0
        )
        self.text_output.pack(side=LEFT, fill=BOTH, expand=True)

        scroll = ttk.Scrollbar(console_frame, command=self.text_output.yview)
        scroll.pack(side=RIGHT, fill=Y)
        self.text_output.config(yscrollcommand=scroll.set)

        self.log("Ready. Please select the MDR document and Project Folder to begin.")

    def log(self, message: str):
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.text_output.insert(END, f"[{timestamp}] {message}\n")
        self.text_output.see(END)
        logging.info(message)

    def select_mdr(self):
        path = filedialog.askopenfilename(
            title="Select MDR (.docx)",
            filetypes=[("Word Document", "*.docx")]
        )
        if path:
            self.mdr_path = path
            self.label_mdr_status.config(text=f"MDR Path: {path}", style="StatusGreen.TLabel")
            self.log(f"Selected MDR file: {path}")

    def select_project_folder(self):
        path = filedialog.askdirectory(
            title="Select Project Folder"
        )
        if path:
            self.project_path = path
            self.label_project_status.config(text=f"Project Folder Path: {path}", style="StatusGreen.TLabel")
            self.log(f"Selected Project Folder: {path}")

    def run_audit(self):
        if not self.mdr_path:
            self.log("ERROR: Please select a valid MDR Word document first.")
            return
        if not self.project_path:
            self.log("ERROR: Please select a valid project directory to audit.")
            return

        try:
            self.log("Step 1: Parsing Master Document Register (MDR)...")
            required_folders, required_files = parse_mdr_docx(self.mdr_path)

            self.log("Step 2: Scanning actual project folder structure...")
            actual_folders, actual_files = scan_project_structure(self.project_path)

            self.log("Step 3: Performing logical gap analysis...")
            nc_list, obs_list, ofi_list, summary = perform_gap_analysis(
                required_folders, required_files, actual_folders, actual_files, self.project_path
            )

            # Output path
            report_dir = os.path.join(self.project_path, "_audit_reports")
            os.makedirs(report_dir, exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            pdf_name = f"ISO_Audit_Report_{timestamp}.pdf"
            pdf_path = os.path.join(report_dir, pdf_name)

            self.log(f"Step 4: Compiling PDF report and styling tables: {pdf_path}")
            generate_pdf_report(
                pdf_path,
                self.project_path,
                self.mdr_path,
                required_folders,
                required_files,
                actual_folders,
                actual_files,
                nc_list,
                obs_list,
                ofi_list,
                summary,
            )

            self.log("SUCCESS: Audit run finished successfully.")
            self.log(f"-> Report saved to: {pdf_path}")
            self.log(f"-> Session logs appended to: {LOG_FILE}")

        except Exception as e:
            logging.exception("Exception occurred during audit execution.")
            self.log(f"CRITICAL ERROR during audit run: {e}")


def main():
    root = Tk()
    root.geometry("900x600")
    # Apply standard theme colors for Windows title bar if supported,
    # otherwise fallback to Tkinter default.
    app = ISOAditorGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
